#!/usr/bin/env python3
"""
Seekurity SIEM 정기점검보고서(Regular Inspection Report) Generator

유지보수 단계의 월간 정기점검 결과 보고서를 .docx 로 생성한다.
서비스 범위 docs/siem-service-scope.md §2.3 (정기 리포트: 시스템 현황 / 탐지 통계 / 개선 권고) 기준.

산출물:
    정기점검보고서_{COMPANY}_{YYYYMMDD}.docx
    - 하드웨어 점검 (서버 구성, 리소스, HW 상태)
    - SIEM 소프트웨어 점검 (서비스/포트, 로그 수집, 인덱스, 탐지 룰, 버전)
    - 보안 이벤트/탐지 통계, 백업 점검, 발견사항/조치, 개선 권고

Usage:
    python inspection_report_manager.py --config examples/demo_inspection_config.json
    python inspection_report_manager.py --company DEMO --date 2026-06-22 --period 2026-06 \
        --inspector "홍길동" --contact "김보안"

config(JSON) 키:
    스칼라 메타: COMPANY_NAME, COMPANY_FULL_NAME, INSPECTION_DATE(YYYY-MM-DD),
        INSPECTION_PERIOD, INSPECTION_TYPE, INSPECTOR, INSPECTOR_ORG,
        CUSTOMER_CONTACT, DOC_VERSION, NEXT_DATE, OVERALL_STATUS, SUMMARY_TEXT
    테이블(리스트, 지정 시 기본 샘플 대체):
        servers, resources, hw_status, siem_services, log_collection,
        index_status, detection_rules, sw_version, detection_stats,
        top_scenarios, backup, findings, recommendations
    그 외 리스트: key_findings(요약 bullet)
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

# Windows 콘솔(cp1252)에서 한글/기호 출력 시 인코딩 오류 방지
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except (AttributeError, ValueError):
    pass

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------------------------------------------------------- #
# 스타일 상수
# ----------------------------------------------------------------------------- #
KR_FONT = "맑은 고딕"
HEADER_FILL = "1F3864"          # 표 머리행 배경 (진한 네이비)
LABEL_FILL = "F2F2F2"           # 라벨 열 배경 (연회색)
BRAND = RGBColor(0x1F, 0x38, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
ORANGE = RGBColor(0xC0, 0x6A, 0x00)
RED = RGBColor(0xC0, 0x39, 0x2B)

STATUS_BAD = ["조치필요", "오류", "위험", "critical", "실패", "중단", "미수신", "불량", "fail", "down"]
STATUS_WARN = ["주의", "경고", "업데이트 가능", "지연", "검토", "진행", "예정", "warn"]
STATUS_GOOD = ["정상", "양호", "running", "green", "최신", "활성", "완료", "성공", "ok", "up"]


# ----------------------------------------------------------------------------- #
# docx 헬퍼
# ----------------------------------------------------------------------------- #
def set_kr(run, name=KR_FONT):
    """런(run)에 한글/영문 폰트를 모두 지정."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), name)


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = KR_FONT
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KR_FONT)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bottom_border(paragraph, color="1F3864", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def status_color(text):
    t = str(text).lower()
    for k in STATUS_BAD:
        if k.lower() in t:
            return RED
    for k in STATUS_WARN:
        if k.lower() in t:
            return ORANGE
    for k in STATUS_GOOD:
        if k.lower() in t:
            return GREEN
    return None


def add_run(paragraph, text, *, bold=False, size=10, color=None):
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    set_kr(run)
    return run


def company_display(m):
    """'정식명 (약칭)' 형태. 정식명이 약칭과 같으면 약칭만."""
    full, name = m.get("COMPANY_FULL_NAME", ""), m["COMPANY_NAME"]
    return f"{full} ({name})" if full and full != name else name


def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, f"{number}. {title}", bold=True, size=14, color=BRAND)
    add_bottom_border(p)
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=11, color=BRAND)
    return p


def add_body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, size=size)
    return p


def add_table(doc, headers, rows, *, widths=None, status_cols=None,
              label_col0=False, center_cols=None, font_size=9.5):
    status_cols = set(status_cols or [])
    center_cols = set(center_cols or [])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i, head in enumerate(headers):
        shade_cell(hdr_cells[i], HEADER_FILL)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, head, bold=True, size=font_size, color=WHITE)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            is_status = i in status_cols
            color = status_color(val) if is_status else None
            run = add_run(p, val, bold=(is_status and color is not None), size=font_size,
                          color=color)
            if label_col0 and i == 0:
                run.bold = True
                shade_cell(cell, LABEL_FILL)
            if i in center_cols or is_status:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = w
    return table


def add_page_number_footer(doc, company, grade=""):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    suffix = f" · {grade}" if grade else ""
    add_run(p, f"Seekurity SIEM 정기점검보고서 · {company}{suffix}   |   ", size=8, color=GREY)
    run = p.add_run()
    set_kr(run)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    for kind, val in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = val
            run._r.append(instr)


# ----------------------------------------------------------------------------- #
# 보고서 본문
# ----------------------------------------------------------------------------- #
def add_cover(doc, m):
    grade = m.get("SECURITY_GRADE", "")
    if grade:
        gp = doc.add_paragraph(); gp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(gp, f"［{grade}］", bold=True, size=11, color=RED)
        spacers = 3
    else:
        spacers = 4
    for _ in range(spacers):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Seekurity SIEM", bold=True, size=18, color=BRAND)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "정기점검보고서", bold=True, size=34, color=BRAND)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Regular Inspection Report", size=13, color=GREY)

    for _ in range(2):
        doc.add_paragraph()

    info = [
        ("고객사", company_display(m)),
        ("점검 유형", m["INSPECTION_TYPE"]),
        ("점검 대상 기간", m["INSPECTION_PERIOD"]),
        ("점검 일자", m["INSPECTION_DATE"]),
        ("점검자", f"{m['INSPECTOR']} / {m['INSPECTOR_ORG']}"),
        ("고객 담당자", m["CUSTOMER_CONTACT"]),
        ("종합 판정", m["OVERALL_STATUS"]),
        ("문서 버전", m["DOC_VERSION"]),
    ]
    table = add_table(
        doc, ["항목", "내용"], info,
        widths=[Cm(4.5), Cm(10.0)],
        status_cols=[],
        label_col0=True,
    )
    # 종합 판정 행만 상태 색상 강조
    judge_cell = table.rows[7].cells[1]
    color = status_color(m["OVERALL_STATUS"])
    if color is not None:
        judge_cell.paragraphs[0].runs[0].font.color.rgb = color
        judge_cell.paragraphs[0].runs[0].bold = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_page_break()


def add_overview(doc, m):
    add_section_heading(doc, 1, "점검 개요")
    add_body(doc,
             f"본 보고서는 {company_display(m)}에 구축된 Seekurity SIEM에 대한 "
             f"{m['INSPECTION_PERIOD']} 정기점검 결과를 정리한 문서이다. 점검은 유지보수 서비스 범위에 따라 "
             "하드웨어 상태, SIEM 소프트웨어 동작, 로그 수집 및 탐지 현황을 대상으로 수행되었다.")
    rows = [
        ("점검 목적", "SIEM 시스템의 안정적 운영 상태 확인 및 잠재 위험 사전 식별"),
        ("점검 범위", "하드웨어(서버/스토리지), SIEM 소프트웨어, 로그 수집·탐지, 백업"),
        ("점검 주기", "월 1회 (정기 방문 점검)"),
        ("점검 방식", "원격 점검 + 정기 방문 (콘솔/CLI 확인, 대시보드 리뷰)"),
        ("판정 기준", "정상 / 주의(개선 권고) / 조치필요(즉시 대응)"),
    ]
    add_table(doc, ["구분", "내용"], rows, widths=[Cm(4.0), Cm(12.5)], label_col0=True)


def add_summary(doc, m):
    add_section_heading(doc, 2, "점검 결과 요약")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "종합 판정: ", bold=True, size=11)
    add_run(p, m["OVERALL_STATUS"], bold=True, size=11, color=status_color(m["OVERALL_STATUS"]) or BRAND)

    add_body(doc, m["SUMMARY_TEXT"])

    add_sub_heading(doc, "주요 발견 사항")
    for item in m["key_findings"]:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, item, size=10)


def add_hardware(doc, m):
    add_section_heading(doc, 3, "하드웨어 점검")

    add_sub_heading(doc, "3.1 서버 구성 현황")
    add_table(
        doc,
        ["호스트명", "역할", "모델", "OS", "CPU / Memory", "판정"],
        m["servers"],
        widths=[Cm(2.8), Cm(2.4), Cm(2.6), Cm(3.0), Cm(3.4), Cm(2.0)],
        status_cols=[5],
    )

    add_sub_heading(doc, "3.2 시스템 리소스 사용 현황")
    add_table(
        doc,
        ["노드", "CPU 평균", "CPU 피크", "Memory", "Disk(/)", "Disk(/data)", "Load", "판정"],
        m["resources"],
        widths=[Cm(2.6), Cm(1.9), Cm(1.9), Cm(1.9), Cm(1.9), Cm(2.1), Cm(1.6), Cm(1.7)],
        status_cols=[7],
        center_cols=[1, 2, 3, 4, 5, 6],
    )

    add_sub_heading(doc, "3.3 하드웨어 상태 점검")
    add_table(
        doc,
        ["점검 항목", "기준", "결과", "판정", "비고"],
        m["hw_status"],
        widths=[Cm(3.6), Cm(3.0), Cm(3.0), Cm(2.0), Cm(4.9)],
        status_cols=[3],
    )


def add_siem_software(doc, m):
    add_section_heading(doc, 4, "SIEM 소프트웨어 점검")

    add_sub_heading(doc, "4.1 서비스 / 프로세스 상태")
    add_table(
        doc,
        ["서비스", "Port", "Protocol", "상태", "자동기동", "비고"],
        m["siem_services"],
        widths=[Cm(4.0), Cm(2.2), Cm(2.0), Cm(2.2), Cm(2.2), Cm(3.9)],
        status_cols=[3],
        center_cols=[1, 2, 4],
    )

    add_sub_heading(doc, "4.2 로그 수집 현황")
    add_table(
        doc,
        ["점검 항목", "값", "판정", "비고"],
        m["log_collection"],
        widths=[Cm(4.5), Cm(3.5), Cm(2.0), Cm(6.5)],
        status_cols=[2],
    )

    add_sub_heading(doc, "4.3 인덱스 / 데이터 상태 (OpenSearch)")
    add_table(
        doc,
        ["점검 항목", "값", "판정", "비고"],
        m["index_status"],
        widths=[Cm(4.5), Cm(3.5), Cm(2.0), Cm(6.5)],
        status_cols=[2],
    )

    add_sub_heading(doc, "4.4 탐지 룰 / 위협 인텔리전스")
    add_table(
        doc,
        ["점검 항목", "값", "판정", "비고"],
        m["detection_rules"],
        widths=[Cm(4.5), Cm(4.5), Cm(2.0), Cm(5.5)],
        status_cols=[2],
    )

    add_sub_heading(doc, "4.5 소프트웨어 버전 및 업데이트")
    add_table(
        doc,
        ["구성요소", "현재 버전", "최신 버전", "판정", "비고"],
        m["sw_version"],
        widths=[Cm(3.6), Cm(2.8), Cm(2.8), Cm(2.0), Cm(5.3)],
        status_cols=[3],
        center_cols=[1, 2],
    )


def add_detection_stats(doc, m):
    add_section_heading(doc, 5, f"보안 이벤트 / 탐지 통계 ({m['INSPECTION_PERIOD']})")

    add_sub_heading(doc, "5.1 심각도별 탐지 현황")
    add_table(
        doc,
        ["심각도", "탐지 건수", "비고"],
        m["detection_stats"],
        widths=[Cm(3.5), Cm(3.5), Cm(9.5)],
        center_cols=[0, 1],
    )

    add_sub_heading(doc, "5.2 주요 탐지 시나리오 Top 5")
    add_table(
        doc,
        ["순위", "탐지 시나리오", "건수", "비고"],
        m["top_scenarios"],
        widths=[Cm(1.6), Cm(7.0), Cm(2.4), Cm(5.5)],
        center_cols=[0, 2],
    )


def add_backup(doc, m):
    add_section_heading(doc, 6, "백업 점검")
    add_table(
        doc,
        ["점검 항목", "주기", "최근 상태", "판정", "비고"],
        m["backup"],
        widths=[Cm(3.6), Cm(2.6), Cm(4.0), Cm(2.0), Cm(4.3)],
        status_cols=[3],
    )


def add_findings(doc, m):
    add_section_heading(doc, 7, "발견 사항 및 조치 내역")
    add_table(
        doc,
        ["No", "구분", "발견 내용", "심각도", "조치 / 계획", "상태"],
        m["findings"],
        widths=[Cm(1.0), Cm(2.4), Cm(5.6), Cm(1.8), Cm(4.0), Cm(1.7)],
        status_cols=[3, 5],
        center_cols=[0],
    )


def add_recommendations(doc, m):
    add_section_heading(doc, 8, "개선 권고 사항")
    add_table(
        doc,
        ["No", "우선순위", "권고 내용", "권장 시기"],
        m["recommendations"],
        widths=[Cm(1.0), Cm(2.4), Cm(9.6), Cm(3.5)],
        status_cols=[1],
        center_cols=[0],
    )


def add_signoff(doc, m):
    add_section_heading(doc, 9, "점검 확인")
    add_body(doc, f"다음 정기점검 예정일: {m['NEXT_DATE']}")
    doc.add_paragraph()
    table = add_table(
        doc,
        ["구분", "소속 / 성명", "확인 (서명)"],
        [
            ("점검자", f"{m['INSPECTOR_ORG']} / {m['INSPECTOR']}", ""),
            ("고객 확인", f"{m['COMPANY_NAME']} / {m['CUSTOMER_CONTACT']}", ""),
        ],
        widths=[Cm(3.0), Cm(7.0), Cm(6.5)],
        center_cols=[],
    )
    # 서명 행 높이 확보
    for r in range(1, 3):
        table.rows[r].height = Cm(1.6)


# ----------------------------------------------------------------------------- #
# 기본(샘플) 데이터
# ----------------------------------------------------------------------------- #
def build_defaults(company, date_iso, period):
    return {
        "COMPANY_NAME": company.upper(),
        "COMPANY_FULL_NAME": f"{company.upper()} 주식회사",
        "INSPECTION_DATE": date_iso,
        "INSPECTION_PERIOD": period,
        "INSPECTION_TYPE": "정기점검 (월간)",
        "INSPECTOR": "홍길동",
        "INSPECTOR_ORG": "Seekurity",
        "CUSTOMER_CONTACT": "보안운영팀",
        "DOC_VERSION": "v1.0",
        "NEXT_DATE": "익월 동일 주차 (협의)",
        "SECURITY_GRADE": "",
        "OVERALL_STATUS": "정상 (주의 3건)",
        "SUMMARY_TEXT": (
            "점검 기간 중 SIEM 시스템은 전반적으로 정상 가동되었으며, 핵심 서비스 및 로그 수집에 "
            "중단은 없었다. 다만 데이터 디스크 사용률 증가, OpenSearch 신규 버전(보안 패치) 출시, "
            "일부 Log Source의 야간 시간대 미수신 등 3건의 주의 항목이 확인되어 개선 권고로 정리하였다."
        ),
        "key_findings": [
            "전 노드 CPU/Memory 사용률 안정 범위 유지, 핵심 서비스 8종 모두 정상 가동.",
            "데이터 디스크 사용률 71% 도달 — 보존정책 조정 또는 증설 검토 필요 (주의).",
            "OpenSearch 2.14.0 신규 버전(보안 패치 포함) 출시 — 업그레이드 협의 필요 (주의).",
            "사용자 DLP 방화벽 02:00~04:00 로그 미수신 구간 확인 — 원인 조치 진행 중 (주의).",
        ],
        # 3.1 서버
        "servers": [
            ("SIEM-MGR-01", "Manager", "Dell R650", "Rocky Linux 9.4", "16 vCPU / 64GB", "정상"),
            ("SIEM-COL-01", "Collector", "Dell R650", "Rocky Linux 9.4", "16 vCPU / 64GB", "정상"),
            ("SIEM-STG-01", "Search/Storage", "Dell R750", "Rocky Linux 9.4", "32 vCPU / 128GB", "정상"),
        ],
        # 3.2 리소스
        "resources": [
            ("SIEM-MGR-01", "28%", "61%", "47%", "38%", "55%", "0.82", "정상"),
            ("SIEM-COL-01", "35%", "72%", "52%", "41%", "63%", "1.10", "정상"),
            ("SIEM-STG-01", "41%", "78%", "66%", "44%", "71%", "1.45", "주의"),
        ],
        # 3.3 HW 상태
        "hw_status": [
            ("전원 공급 장치(PSU)", "이중화 정상", "이중화 정상", "정상", "PSU 2기 모두 가동"),
            ("시스템 온도", "< 35°C", "31°C", "정상", "iDRAC 기준"),
            ("냉각 팬(Fan)", "전체 정상", "전체 정상", "정상", "-"),
            ("RAID 상태", "Optimal", "RAID-10 Optimal", "정상", "Storage 노드 데이터 볼륨"),
            ("물리 디스크", "불량 0", "0개", "정상", "Predictive Failure 0"),
            ("네트워크 인터페이스", "Link Up (10G)", "Link Up", "정상", "Bonding 정상"),
            ("시스템 Uptime", "-", "87일", "정상", "계획 외 재부팅 없음"),
            ("하드웨어 경보(iDRAC)", "경보 없음", "0건", "정상", "-"),
        ],
        # 4.1 서비스 (포트: README 표준 포트 기준)
        "siem_services": [
            ("Nginx (Web Console)", "443 / 80", "TCP", "Running", "활성", "HTTPS 정상"),
            ("SS-Syslog-Receiver", "514", "UDP", "Running", "활성", "수신 정상"),
            ("SS-API", "23001", "TCP", "Running", "활성", "-"),
            ("SS-Console", "23002", "TCP", "Running", "활성", "-"),
            ("OpenSearch", "19200 / 19300", "TCP", "Running", "활성", "API/Transport"),
            ("PostgreSQL", "15432", "TCP", "Running", "활성", "-"),
            ("Kafka", "19092", "TCP", "Running", "활성", "-"),
            ("Zookeeper", "12181", "TCP", "Running", "활성", "-"),
        ],
        # 4.2 로그 수집
        "log_collection": [
            ("연동 Log Source 총계", "67 식", "정상", "-"),
            ("정상 수신", "66 식", "정상", "-"),
            ("수신 지연 / 미수신", "1 식", "주의", "사용자 DLP 방화벽 야간 미수신"),
            ("평균 EPS", "3,200 EPS", "정상", "월 평균"),
            ("피크 EPS", "8,700 EPS", "정상", "용량 임계 대비 여유"),
            ("일일 평균 수집량", "약 142 GB/day", "정상", "-"),
        ],
        # 4.3 인덱스
        "index_status": [
            ("Cluster Health", "green", "정상", "-"),
            ("노드 수", "3", "정상", "-"),
            ("활성 Shard", "412", "정상", "-"),
            ("Unassigned Shard", "0", "정상", "-"),
            ("인덱스 수", "248", "정상", "-"),
            ("총 문서 수", "약 18.4억 건", "정상", "-"),
            ("데이터 디스크 사용률", "71%", "주의", "80% 도달 전 보존정책/증설 검토"),
            ("보존 정책", "Hot 90일 / Archive 365일", "정상", "정책대로 적용 중"),
        ],
        # 4.4 탐지 룰 / TI
        "detection_rules": [
            ("활성 탐지 룰", "56 개", "정상", "-"),
            ("룰 동작 점검", "정상", "정상", "샘플 이벤트 매칭 확인"),
            ("미동작 / 오류 룰", "0 개", "정상", "-"),
            ("위협 인텔리전스 Feed", "최근 갱신 2026-06-21", "정상", "일 1회 자동 갱신"),
            ("등록 IOC", "615 건", "정상", "-"),
        ],
        # 4.5 SW 버전
        "sw_version": [
            ("Seekurity SIEM", "v1.0", "v1.0", "최신", "-"),
            ("OpenSearch", "2.13.0", "2.14.0", "업데이트 가능", "보안 패치 포함, 업그레이드 협의"),
            ("PostgreSQL", "15.6", "15.7", "업데이트 가능", "마이너 패치"),
            ("OS (Rocky Linux)", "9.4", "9.4", "최신", "보안 패치 적용 완료"),
        ],
        # 5.1 탐지 통계
        "detection_stats": [
            ("Critical", "2", "즉시 분석/대응 완료"),
            ("High", "18", "전건 확인, 오탐 4건 튜닝"),
            ("Medium", "134", "주기 검토"),
            ("Low", "1,247", "참고 모니터링"),
            ("총계", "1,401", "-"),
        ],
        # 5.2 Top 시나리오
        "top_scenarios": [
            ("1", "무차별 대입 로그인 시도 (Brute Force)", "312", "차단/계정 보호 확인"),
            ("2", "비인가 포트 스캔", "187", "외부 출발지 다수"),
            ("3", "악성 IP 통신 시도 (TI 매칭)", "96", "방화벽 차단 연계"),
            ("4", "Firewall Deny 급증", "64", "특정 시간대 집중"),
            ("5", "권한 상승 시도", "11", "관리계정 모니터링"),
        ],
        # 6 백업
        "backup": [
            ("설정 백업", "일 1회", "최근 성공 2026-06-22 02:00", "정상", "자동 백업"),
            ("데이터 스냅샷", "주 1회", "최근 성공 2026-06-21", "정상", "OpenSearch snapshot"),
            ("백업 저장소 사용률", "-", "58%", "정상", "여유 확보"),
            ("복구 테스트", "분기 1회", "2026-04 수행", "정상", "정상 복구 확인"),
        ],
        # 7 발견사항
        "findings": [
            ("1", "하드웨어", "데이터 디스크 사용률 71% 도달 (증가 추세)", "주의",
             "보존정책 조정 또는 디스크 증설 검토", "검토중"),
            ("2", "SIEM SW", "OpenSearch 2.14.0(보안 패치) 신규 출시", "주의",
             "다음 정기점검 시 업그레이드 협의", "예정"),
            ("3", "로그 수집", "사용자 DLP 방화벽 02:00~04:00 로그 미수신", "주의",
             "장비 측 Syslog 정책/시간대 점검 요청", "진행중"),
        ],
        # 8 권고
        "recommendations": [
            ("1", "높음", "데이터 디스크 증설 또는 보존 기간 조정 (현 71%, 증가 추세)", "2026 Q3 내"),
            ("2", "중간", "OpenSearch 2.14.0 업그레이드 (보안 패치 적용)", "다음 정기점검"),
            ("3", "중간", "사용자 DLP 방화벽 로그 미수신 구간 근본 원인 조치", "2주 내"),
            ("4", "낮음", "Low 등급 이벤트 다량 발생 — 탐지 룰 임계치 튜닝 검토", "상시"),
        ],
    }


def merge_config(defaults, config):
    merged = dict(defaults)
    for key, value in config.items():
        if isinstance(value, list) and key in merged:
            merged[key] = [tuple(v) if isinstance(v, list) else v for v in value]
        else:
            merged[key] = value
    return merged


# ----------------------------------------------------------------------------- #
# 생성
# ----------------------------------------------------------------------------- #
def generate_report(m, output_dir=None):
    doc = Document()
    set_base_style(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 코어 속성
    doc.core_properties.title = f"Seekurity SIEM 정기점검보고서 - {m['COMPANY_NAME']}"
    doc.core_properties.author = m["INSPECTOR_ORG"]
    doc.core_properties.subject = f"{m['INSPECTION_PERIOD']} 정기점검"

    add_page_number_footer(doc, m["COMPANY_NAME"], m.get("SECURITY_GRADE", ""))

    add_cover(doc, m)
    add_overview(doc, m)
    add_summary(doc, m)
    add_hardware(doc, m)
    add_siem_software(doc, m)
    add_detection_stats(doc, m)
    add_backup(doc, m)
    add_findings(doc, m)
    add_recommendations(doc, m)
    add_signoff(doc, m)

    # 출력 경로 결정
    scripts_dir = Path(__file__).parent
    siem_root = scripts_dir.parent
    if output_dir:
        out_dir = Path(output_dir)
    else:
        proj = siem_root / "projects" / m["COMPANY_NAME"].lower()
        out_dir = (proj / "reports") if proj.exists() else (siem_root / "output")
    out_dir.mkdir(parents=True, exist_ok=True)

    date_compact = datetime.strptime(m["INSPECTION_DATE"], "%Y-%m-%d").strftime("%Y%m%d")
    out_file = out_dir / f"정기점검보고서_{m['COMPANY_NAME'].upper()}_{date_compact}.docx"
    doc.save(out_file)
    return out_file


def main():
    parser = argparse.ArgumentParser(
        description="Seekurity SIEM 정기점검보고서(.docx) 생성기",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--config", type=str, help="JSON config 파일 경로")
    parser.add_argument("--company", type=str, help="고객사명 (예: KOVAN)")
    parser.add_argument("--company-full", type=str, help="고객사 정식 명칭")
    parser.add_argument("--date", type=str, help="점검 일자 (YYYY-MM-DD, 기본: 오늘)")
    parser.add_argument("--period", type=str, help="점검 대상 기간 (예: 2026-06)")
    parser.add_argument("--inspector", type=str, help="점검자명")
    parser.add_argument("--contact", type=str, help="고객 담당자")
    parser.add_argument("--output-dir", type=str, help="출력 디렉터리 (미지정 시 자동)")
    args = parser.parse_args()

    config = {}
    if args.config:
        with open(args.config, "r", encoding="utf-8") as f:
            config = json.load(f)

    company = (args.company or config.get("COMPANY_NAME") or "DEMO")
    date_iso = (args.date or config.get("INSPECTION_DATE") or datetime.now().strftime("%Y-%m-%d"))
    # 기간 미지정 시 점검월 기준
    period = (args.period or config.get("INSPECTION_PERIOD")
              or datetime.strptime(date_iso, "%Y-%m-%d").strftime("%Y-%m"))

    data = build_defaults(company, date_iso, period)

    cli_overrides = {
        "COMPANY_NAME": args.company,
        "COMPANY_FULL_NAME": args.company_full,
        "INSPECTION_DATE": args.date,
        "INSPECTION_PERIOD": args.period,
        "INSPECTOR": args.inspector,
        "CUSTOMER_CONTACT": args.contact,
    }
    config = {**config, **{k: v for k, v in cli_overrides.items() if v}}

    data = merge_config(data, config)

    try:
        out_file = generate_report(data, args.output_dir)
    except Exception as exc:  # noqa: BLE001
        print(f"Error: {exc}")
        raise

    print(f"✓ 정기점검보고서 생성 완료: {out_file}")
    print(f"  고객사: {data['COMPANY_FULL_NAME']} ({data['COMPANY_NAME']})")
    print(f"  점검 기간: {data['INSPECTION_PERIOD']}  /  점검 일자: {data['INSPECTION_DATE']}")
    print(f"  종합 판정: {data['OVERALL_STATUS']}")


if __name__ == "__main__":
    main()
